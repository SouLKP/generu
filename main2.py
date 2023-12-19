import streamlit as st

st.title('Resume Enhancer')


title_query = '''you have resume data from that extract the name of the person,experience,and job title.
In following format :

Name of the person - job title ([number] years of Experience)'''

summary_query = '''extract only summary detail from the content.  
below some prefenied word format are there, so output needed in proper to make sure:

Title (like: Summary)
[points]

***output should be in proper format***                          

Summary
7+ years of experience as a Python developer...
...only all summary related point needed
...                            
'''

skill_query = '''extract the only skills from the content.  
below some predefined word format are there, so output needed in below proper to make sure:

Title (like :Skills)
[skill name or category]

Examples : [Programming Languages]: [Python, Angular, ...]
 
***output should be in proper format***

Skills

Programming Languages: Python, Angular, JavaScript, SQL
Framework: Flask
Tools: Kubernetes, Jupiter Notebooks, Kafta, Docker
...for all skill

'''

# # project_query = '''extract the projects details with from the content.  
# # below some predefined word format are there, so output needed in below proper to make sure:

# # [project name] 
# # Role : [role] 
# # [description of the project]
 
# # ***output should be in proper format***

# # 1. Mining Site Report generation
# # Role: Team Lead 
# # description : A definite report of .... .

# # ...for all projects
# # '''

project_query = '''you have resume details from that you have to extract the "project" related details.  
below some predefined word format are there, so output needed in below proper to make sure:

[project name] 
Role : [role] 
[description of the project]
 
***output should be in proper format***
*** project name, role , description **** if not then create detail based on job title from content

Example :

1. Mining Site Report generation
Role: Team Lead 
description : A definite report of ...

...
...

'''


role_query = '''

Roles and Responsibilities
List 3-5 points stating key roles and responsibilities related to the [job title] position, based on the background details provided.

Example point:
 Managed software development of ...

If any details are missing from the provided content, create Roles and Responsibilities infomation based on that person background to complete the requested resume format. Focus on tailoring details to a specific job title.

***output should be in proper format***

Roles and Responsibilities
-point
-point
...

'''

address = '''Signature 1, 1007-1010, 10th floor,
Sarkhej - Gandhinagar Highway,Makarba,
Ahmedabad, Gujarat 380051'''

from llama_index.llms import OpenAI
from llama_index import VectorStoreIndex, SimpleDirectoryReader, ServiceContext

from dotenv import load_dotenv
load_dotenv()
import os
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
def first_query_engine(directory_path):
    llm = OpenAI(temperature=0.1, model="gpt-3.5-turbo-16k", max_tokens=5000, timeout= 300) 
    service_context = ServiceContext.from_defaults(llm=llm) 
    documents = SimpleDirectoryReader(directory_path).load_data()  
    index1 = VectorStoreIndex.from_documents(documents, service_context=service_context) 
    query_engine1 = index1.as_query_engine() 
    return query_engine1

import shutil
current_path = os.getcwd()
st.write(current_path,"current path")
path_1 = os.path.join(current_path, 'resume_enhancer1')
path_2 = os.path.join(current_path,'output')
os.makedirs(path_1, exist_ok=True)
os.makedirs(path_2,exist_ok=True)

if os.path.exists(path_1):
    st.write(path_1,"resume path")
if os.path.exists(path_2):
    st.write(path_2,"output file")

# full_path = os.path.join(save_path, resume1.name)

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

from docx.shared import Cm
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# # def create_summary_docx(title,address,summary,skill,role,projects_string,file_path): 
# #         role = role.replace('-','')
# #         summary = summary.replace('●','')
# #         doc =Document()
        
# #         section = doc.sections[0]
# #         header = section.header
# #         header.add_paragraph().add_run().add_picture('/home/webclues/Documents/clog.png', width=Cm(5.44), height=Cm(0.69))
# #         doc.add_paragraph().add_run().add_break()

# #         #started title 
# #         title_paragraph = doc.add_paragraph(title) 
# #         title_paragraph.paragraph_format.line_spacing = 0.8
# #         title_run = title_paragraph.runs[0]
# #         title_run.bold = True
# #         title_run.font.size = Pt(15)
# #         title_run.font.color.rgb = RGBColor(102, 101, 101) 
# #         #ended title

# #         #line
# #         def add_horizontal_line(paragraph, rgb_color, thickness_pt):
# #             pBdr = OxmlElement('w:pBdr')
# #             bottom_border = OxmlElement('w:bottom')
# #             bottom_border.set(qn('w:val'), 'single')
# #             bottom_border.set(qn('w:color'), rgb_color)
# #             pBdr.append(bottom_border)
# #             paragraph._element.get_or_add_pPr().append(pBdr)

# #         line_paragraph = doc.add_paragraph()
# #         add_horizontal_line(line_paragraph, '585353', 6)  
# #         doc.add_paragraph().add_run('').add_break()

# #         #address
# #         title_paragraph = doc.add_paragraph("Codezeros") 
# #         title_run = title_paragraph.runs[0]
# #         title_run.bold = False
# #         title_run.font.size = Pt(18.5)
# #         title_run.font.color.rgb = RGBColor(204,64,37) 
# #         lines = address.split('\n')
# #         for i in lines:
# #             title_paragraph = doc.add_paragraph(i)
# #             title_paragraph.paragraph_format.line_spacing = 0.3
# #             title_run = title_paragraph.runs[0]
# #             title_run.bold = False
# #             title_run.font.size = Pt(11)
# #             title_run.font.color.rgb = RGBColor(102, 102, 102)
# #         doc.add_paragraph().add_run('').add_break()

# #         #summray
# #         lines = summary.split('\n')
# #         title1 = doc.add_paragraph('Summary ')
# #         run = title1.runs[0]
# #         run.bold = True
# #         run.font.size = Pt(16)
# #         run.font.color.rgb = RGBColor(204,85,0)   
# #         title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# #         lines = lines[1:]  
# #         doc.add_paragraph().add_run('').add_break()
 
# #         for point in lines:
# #             if point.strip() != '':
# #                 p = doc.add_paragraph(point, style='List Bullet')
# #                 p.add_run(point).font.size = Pt(12)
# #                 p.paragraph_format.line_spacing = 2

# #         #skills
# #         lines = skill.split('\n') 
# #         title1 = doc.add_paragraph(lines[0])
# #         run = title1.runs[0]
# #         run.bold = True
# #         run.font.size = Pt(16)
# #         run.font.color.rgb = RGBColor(204,85,0)   
# #         title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# #         lines = lines[1:]   
# #         doc.add_paragraph().add_run('').add_break()

# #         for line in lines:
# #             if line.strip() != '':
# #                 p = doc.add_paragraph(line, style='List Bullet')
# #                 p.add_run(line).font.size = Pt(12)
# #                 p.paragraph_format.line_spacing = 2
        
# #         #role
# #         lines = role.split('\n')
# #         title1 = doc.add_paragraph(lines[0])
# #         run = title1.runs[0]
# #         run.bold = True
# #         run.font.size = Pt(16)
# #         run.font.color.rgb = RGBColor(204,85,0)   
# #         title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# #         lines = lines[1:]  
# #         doc.add_paragraph().add_run('').add_break() 
# #         for point in lines:
# #             if point.strip() != ' ' or '' or '-' :
# #                 p = doc.add_paragraph(point, style='List Bullet')
# #                 p.add_run().font.size = Pt(12)
# #                 p.paragraph_format.line_spacing = 2

# #         #projects
# #         def add_project(doc, project_name, role, description):
# #             # Add project name
# #             project_heading = doc.add_paragraph(f'{project_name}')
# #             project_heading.runs[0].bold = True
# #             project_heading.runs[0].font.size = Pt(12)
# #             project_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# #             # Add role
# #             role_heading = doc.add_paragraph('Role:')
# #             role_heading.runs[0].bold = True
# #             role_paragraph = doc.add_paragraph(role)
# #             role_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# #             # Add description
# #             description_heading = doc.add_paragraph('Description:')
# #             description_heading.runs[0].bold = True
# #             description_paragraph = doc.add_paragraph(description)
# #             description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
 
# #             doc.add_paragraph().add_run().add_break()
    
# #         def parse_projects(projects_string):
# #             # Split the string into individual projects
# #             projects = projects_string.strip().split('\n\n')
# #             for project in projects:
# #                 # Split the project details
# #                 title, rest = project.split('\n', 1)
# #                 role, description = rest.split('\n', 1)
# #                 yield title.strip(), role.split(': ', 1)[1].strip(), description.strip()

# #         def create_document_with_projects(projects_string, file_path):
# #             title = doc.add_paragraph('Projects')
# #             run = title.runs[0]
# #             run.bold = True
# #             run.font.size = Pt(16)
# #             run.font.color.rgb = RGBColor(204,85,0)
# #             title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# #             doc.add_paragraph().add_run().add_break()

# #             for project_name, role, description in parse_projects(projects_string):
# #                 add_project(doc, project_name, role, description)

# #             doc.save(file_path)
            
# #         create_document_with_projects(projects_string,file_path)
# #         return True

def create_summary_docx(title,address,summary,skill,role,projects_string,file_path,resume_format): 
        if resume_format == 'web':
            filename = '/home/webclues/Documents/webclues_logo.png'
            width = Cm(7.30)
            height = Cm(1.25)
            #title color
            #line color
            #address_color
            address_name = 'Webclues'
            #address,summary,skiil,role color
            color = RGBColor(0,0,254)
        else:
            filename = '/home/webclues/Documents/clog.png'
            width = Cm(5.44)
            height = Cm(0.69)
            address_name = 'CodeZeros'
            color = RGBColor(204,64,37)
        
        doc =Document()

        section = doc.sections[0]
        header = section.header
        header.add_paragraph().add_run().add_picture(filename, width=width, height=height)
        doc.add_paragraph().add_run().add_break()

        #started title 
        title_paragraph = doc.add_paragraph(title) 
        title_paragraph.paragraph_format.line_spacing = 0.8
        title_run = title_paragraph.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(15)
        title_run.font.color.rgb = RGBColor(102,101,101) 
        #ended title

        #line
        def add_horizontal_line(paragraph, rgb_color, thickness_pt):
            pBdr = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:color'), rgb_color)
            pBdr.append(bottom_border)
            paragraph._element.get_or_add_pPr().append(pBdr)

        line_paragraph = doc.add_paragraph()
        add_horizontal_line(line_paragraph, '585353', 6)  
        doc.add_paragraph().add_run('').add_break()

        #address
        title_paragraph = doc.add_paragraph(address_name) 
        title_run = title_paragraph.runs[0]
        title_run.bold = False
        title_run.font.size = Pt(18.5)
        title_run.font.color.rgb = color 
        lines = address.split('\n')
        for i in lines:
            title_paragraph = doc.add_paragraph(i)
            title_paragraph.paragraph_format.line_spacing = 0.2 #0.3
            title_run = title_paragraph.runs[0]
            title_run.bold = False
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = RGBColor(102, 102, 102)
        doc.add_paragraph().add_run('').add_break()

        #summray
        summary = summary.replace('-','')
        lines = summary.split('\n')
        title1 = doc.add_paragraph('Summary ')
        run = title1.runs[0]
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = color   
        title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        lines = lines[1:]  
        doc.add_paragraph().add_run('').add_break()
 
        for point in lines:
            if point.strip() != '':
                p = doc.add_paragraph(point, style='List Bullet')
                p.add_run().font.size = Pt(12)
                p.paragraph_format.line_spacing = 2

        #skills
        title1 = doc.add_paragraph("Skills")
        run = title1.runs[0]
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = color 
        title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # for line in lines:
        #     if line.strip() != '':
        #         p = doc.add_paragraph(line, style='List Bullet')
        #         p.add_run(line).font.size = Pt(12)
        #         p.paragraph_format.line_spacing = 2
                
        #skill
        data = []
        lines = skill.strip().split('\n')[2:]  

        for line in lines:
            if line != '' or ' ':
                if '|' in line:
                    main_category, sub_category_str = line.split('|')
                    main_category = main_category.strip()
                    sub_categories = [sub.strip() for sub in sub_category_str.split(',')]
                    data.append((main_category, sub_categories))

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Main Category'
        header_cells[1].text = 'Sub Category'

        for main_category, sub_categories in data:
            row_cells = table.add_row().cells
            row_cells[0].text = main_category
            row_cells[1].text = sub_categories[0]

            # For each subsequent subcategory, add a new row and populate the second cell
            for sub_category in sub_categories[1:]:
                row_cells = table.add_row().cells
                row_cells[1].text = sub_category

            # Merge cells for the main category
            if len(sub_categories) > 1:  # More than one subcategory, so we merge
                a = table.cell(len(table.rows) - len(sub_categories), 0)
                b = table.cell(len(table.rows) - 1, 0)
                a.merge(b)
        doc.add_paragraph().add_run().add_break()


        
        #role
        summary = summary.replace('-','')
        lines = role.split('\n')
        title1 = doc.add_paragraph(lines[0])
        run = title1.runs[0]
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = color   
        title1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        lines = lines[1:]  
        for point in lines:
            if point.strip() != ' ' or '' or '-' :
                p = doc.add_paragraph(point, style='List Bullet')
                p.add_run(point).font.size = Pt(12)
                p.paragraph_format.line_spacing = 2
        doc.add_paragraph().add_run('').add_break()

        #projects
        def add_project(doc, project_name, role, description):
            # Add project name
            project_heading = doc.add_paragraph(f'{project_name}')
            project_heading.runs[0].bold = True
            project_heading.runs[0].font.size = Pt(12)
            project_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Add role
            role_heading = doc.add_paragraph('Role:')
            role_heading.runs[0].bold = True
            role_paragraph = doc.add_paragraph(role)
            role_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Add description
            description_heading = doc.add_paragraph('Description:')
            description_heading.runs[0].bold = True
            description_paragraph = doc.add_paragraph(description)
            description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
 
            doc.add_paragraph().add_run().add_break()
    
        def parse_projects(projects_string):
            # Split the string into individual projects
            projects = projects_string.strip().split('\n\n')
            for project in projects:
                # Split the project details
                title, rest = project.split('\n', 1)
                role, description = rest.split('\n', 1)
                yield title.strip(), role.split(': ', 1)[1].strip(), description.strip()

        def create_document_with_projects(projects_string, file_path):
            title = doc.add_paragraph('Projects')
            run = title.runs[0]
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = color
            title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph().add_run().add_break()

            for project_name, role, description in parse_projects(projects_string):
                add_project(doc, project_name, role, description)

            doc.save(file_path)
            
        create_document_with_projects(projects_string,file_path)
        return True

import subprocess

def convert_to_pdf(input_file, output_dir):
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', input_file, '--outdir', output_dir])
    return True

# input_file = os.path.join('/home/webclues/Music/HR_Gen_3/dataset1/output','final_resume.docx')
# output_dir = os.path.join('/home/webclues/Music/HR_Gen_3/dataset1/output')  

if 'allow1' not in st.session_state:
    st.session_state.allow1 = False

if 'allow2' not in st.session_state:
    st.session_state.allow2 = False

if 'directory_path' not in st.session_state:
    st.session_state.directory_path = False

# resume1 = st.file_uploader("Choose Vendor's Resume", type=['docx', 'pdf'], key='p1')
# radio = st.radio(label="Select Resume Format",options=['code','web'],horizontal=False)


def updated_skill(skill_text):
    from openai import OpenAI
    client = OpenAI()
    response = client.chat.completions.create(model="gpt-4",messages=[{"role": "system", "content": "You are a intelligent coder assistant."},
    {"role": "user", "content": f"understand this skill details : {skill_text}.understand skills details and based on that make two columns for main category and right side subcategory needed.output two column sperate wtih '|' "}])
    return response.choices[0].message.content

def engine_query(query1,engine):
    response = engine.query(query1)
    return response.response



with st.sidebar.form(key='pdf_upload_form', clear_on_submit=True):
    resume1 = st.file_uploader("Choose Vendor's Resume", type=['docx', 'pdf'], key='p1')
    radio = st.radio(label="Select Resume Format", options=['code', 'web'], horizontal=True)
    submit_button = st.form_submit_button(label='Generate Resume', on_click=None)

# if st.button("Generate Resume",type='primary'):
if submit_button:    
    if resume1 is not None:
        directory_path = os.path.join(path_1,resume1.name)
        os.makedirs(directory_path,exist_ok=True)
        if os.path.exists(directory_path):
            st.write(directory_path,"directory path done")
        file_path = os.path.join(directory_path,resume1.name)
        # st.session_state.directory_path = directory_path
        st.write(file_path,"file path done")
        with open(file_path, "wb") as f:
            f.write(resume1.getbuffer())
            f.close()
        if os.path.exists(file_path):
            st.write("file stored")
        st.write("content",'data!!!!!!!!!!!!!!!!!!')
        engine = first_query_engine(directory_path)
        # q1 = engine_query(title_query,engine)
        # st.write(q1)

        st.session_state.allow1 = True
        st.session_state.allow2 = True 

#         # Example usage
#         file_path = '/home/webclues/Music/HR_Gen_3/dataset1/output/final_resume.docx'
#         dataset_path = '/home/webclues/Music/HR_Gen_3/dataset1/resume_enhancer2'
        input_file = os.path.join(path_2,'final_resume.docx')
        output_dir = os.path.join(path_2)  

        title = engine_query(title_query,engine)
        st.write(title)
        summary = engine_query(summary_query,engine)
        st.write(summary)
        skill_text = engine_query(skill_query,engine)
        st.write(skill_text)
        skill = updated_skill(skill_text)
        st.write(skill)
        role = engine_query(role_query,engine)
        st.write(role)
        projects_string = engine_query(project_query,engine)
        st.write(projects_string)

        role = role.replace('-','')
        summary = summary.replace('●','')
          
        create_summary_docx(title, address, summary, skill, role, projects_string,file_path,resume_format = radio)
        convert_to_pdf(input_file, output_dir)
        st.session_state.directory_path = True
    

file_path1 = os.path.join(path_2,'final_resume.pdf')
file_path2 = os.path.join(path_2,'final_resume.docx')
 
if st.session_state.directory_path == True:
    if st.session_state.allow1 == True:
        with open(file_path1, "rb") as file:
            if st.download_button(label="Download Document PDF", data=file, file_name="document.pdf", mime="application/pdf"):
                st.session_state.allow1 = False

if st.session_state.directory_path == True:
    if st.session_state.allow2 == True:
        with open(file_path2, "rb") as file:
            if st.download_button(label="Download Document DOCX", data=file, file_name="document.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                st.session_state.allow2 = False
 
   
